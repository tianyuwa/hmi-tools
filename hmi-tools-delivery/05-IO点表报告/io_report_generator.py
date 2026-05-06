#!/usr/bin/env python3
"""
IO 点表自动生成报告 v2.0
========================
专业 HMI/PLC IO 点表报告生成工具。
支持自动分类（DI/DO/AI/AO）、地址占用统计、槽位汇总、余量分析。
导出: Word (.docx) · PDF · Markdown · CSV

使用:
    python io_report_generator.py --example                    # 生成示例报告
    python io_report_generator.py -i tags.csv -o output_dir    # 从标签文件生成
    python io_report_generator.py -i tags.csv --skip-word      # 跳过 Word 导出

依赖 (可选):
    pip install python-docx    → .docx 格式
    pip install fpdf2          → .pdf 格式

作者: HMI Toolbox
"""

import csv
import json
import os
import sys
import argparse
import re
from datetime import datetime
from collections import defaultdict

# ═══════════════════════════════════════════════════════════════
# 依赖检测
# ═══════════════════════════════════════════════════════════════

HAS_DOCX = False
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    pass

HAS_FPDF = False
try:
    import warnings
    warnings.filterwarnings('ignore', message='.*NOT subset.*')
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    pass

# ═══════════════════════════════════════════════════════════════
# 常量
# ═══════════════════════════════════════════════════════════════

IO_LABELS = {
    'DI': '数字量输入',
    'DO': '数字量输出',
    'AI': '模拟量输入',
    'AO': '模拟量输出',
    'Internal': '内部变量',
    'Other': '其他',
}

IO_ORDER = ['DI', 'DO', 'AI', 'AO', 'Internal', 'Other']

# 专业配色
IO_COLORS_HEX = {
    'DI': '#2196F3',
    'DO': '#FF9800',
    'AI': '#4CAF50',
    'AO': '#F44336',
}

IO_COLORS_RGB = {
    'DI': (33, 150, 243),
    'DO': (255, 152, 0),
    'AI': (76, 175, 80),
    'AO': (244, 67, 54),
}

# ═══════════════════════════════════════════════════════════════
# 数据模型
# ═══════════════════════════════════════════════════════════════

class IOPoint:
    def __init__(self, name, address, data_type='BOOL', comment='',
                 io_type='DI', card_slot='', terminal='', cable=''):
        self.name = name
        self.address = address
        self.data_type = data_type
        self.comment = comment
        self.io_type = io_type
        self.card_slot = card_slot
        self.terminal = terminal
        self.cable = cable

# ═══════════════════════════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════════════════════════

def classify_io(address, data_type):
    """根据地址和数据类型自动分类 IO 类型"""
    addr = address.upper().strip()
    dt = data_type.upper().strip()

    if re.match(r'%?I[A-Z]*\d', addr) or re.match(r'[EI]\d', addr):
        return 'DI' if dt in ('BOOL', 'BIT', '') else 'AI'
    if re.match(r'%?Q[A-Z]*\d', addr) or re.match(r'[A]\d', addr):
        return 'DO' if dt in ('BOOL', 'BIT', '') else 'AO'
    if re.match(r'%?AI', addr) or re.match(r'%?IW', addr):
        return 'AI'
    if re.match(r'%?AQ', addr) or re.match(r'%?QW', addr):
        return 'AO'
    if addr.startswith('X'):
        return 'DI'
    if addr.startswith('Y'):
        return 'DO'
    if addr.startswith(('M', 'DB', 'V')):
        return 'Internal'
    return 'Other'


def auto_detect_delimiter(file_path):
    """自动检测 CSV 分隔符"""
    with open(file_path, 'r', encoding='utf-8-sig') as f:
        header = f.readline()
    if ',' in header and header.count(',') >= header.count(';'):
        return ','
    return ';'


def find_chinese_font():
    """查找系统中的中文字体"""
    candidates = [
        r'C:\Windows\Fonts\msyh.ttc',
        r'C:\Windows\Fonts\msyhbd.ttc',
        r'C:\Windows\Fonts\simsun.ttc',
        r'C:\Windows\Fonts\deng.ttf',
        r'C:\Windows\Fonts\yahei.ttf',
        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        '/System/Library/Fonts/STHeiti Light.ttc',
        '/System/Library/Fonts/PingFang.ttc',
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


def _set_word_cell_shading(cell, hex_color):
    """设置 Word 单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _style_word_table(table, header_texts, data_rows,
                      col_widths=None, header_bg='0F172A', alt_bg='F5F7FA',
                      font_size=9):
    """通用 Word 表格格式化"""
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                if i < len(row.cells):
                    row.cells[i].width = w

    # 表头
    for i, h in enumerate(header_texts):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_word_cell_shading(cell, header_bg)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = '微软雅黑'

    # 数据行
    for row_idx in range(1, len(table.rows)):
        for ci in range(len(table.rows[row_idx].cells)):
            if (row_idx - 1) % 2 == 0:
                _set_word_cell_shading(table.rows[row_idx].cells[ci], alt_bg)
            for para in table.rows[row_idx].cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = '微软雅黑'

    return table


def _make_word_header(doc, text, level=1):
    """创建带格式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
    return h


def _make_word_para(doc, text, size=10, color=None, bold=False, italic=False, font_name='微软雅黑'):
    """创建带格式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return p


# ═══════════════════════════════════════════════════════════════
# 报告生成器主类
# ═══════════════════════════════════════════════════════════════

class IOReportGenerator:

    @staticmethod
    def read_tags(file_path):
        """读取标签文件，自动检测列名映射"""
        if not os.path.exists(file_path):
            return []

        delimiter = auto_detect_delimiter(file_path)

        with open(file_path, 'r', encoding='utf-8-sig') as f:
            reader = csv.DictReader(f, delimiter=delimiter)
            if not reader.fieldnames:
                return []

            # 智能列名映射
            col_map = {}
            for fn in reader.fieldnames:
                fu = fn.upper().strip()
                if fu in ('变量名', 'NAME', '标记', 'TAG', 'SYMBOL', 'TAG NAME', 'TAGNAME'):
                    col_map['name'] = fn
                elif fu in ('地址', 'ADDRESS', 'ADDR', 'PLC地址', 'PLC_ADDR', 'PLCADDR'):
                    col_map['address'] = fn
                elif fu in ('数据类型', 'DATATYPE', 'DATA_TYPE', 'TYPE'):
                    col_map['data_type'] = fn
                elif fu in ('备注', '注释', 'COMMENT', '说明', 'DESCRIPTION', 'DESC'):
                    col_map['comment'] = fn
                elif fu in ('槽位', 'SLOT', 'CARD_SLOT', 'CARDSLOT'):
                    col_map['card_slot'] = fn
                elif fu in ('端子', 'TERMINAL', 'TERM'):
                    col_map['terminal'] = fn
                elif fu in ('电缆', 'CABLE', 'WIRE'):
                    col_map['cable'] = fn

            points = []
            for row in reader:
                name = row.get(col_map.get('name', ''), '').strip()
                addr = row.get(col_map.get('address', ''), '').strip()
                if not name or not addr:
                    continue
                dtype = row.get(col_map.get('data_type', ''), 'BOOL').strip().upper()
                comment = row.get(col_map.get('comment', ''), '').strip()
                card = row.get(col_map.get('card_slot', ''), '').strip()
                term = row.get(col_map.get('terminal', ''), '').strip()
                cable = row.get(col_map.get('cable', ''), '').strip()
                io_type = classify_io(addr, dtype)
                points.append(IOPoint(name, addr, dtype, comment, io_type, card, term, cable))

        return points

    # ── 内部辅助 ──

    @staticmethod
    def _classify(points):
        c = defaultdict(list)
        for p in points:
            c[p.io_type].append(p)
        return c

    @staticmethod
    def _analyze_addresses(points):
        """地址区域占用分析"""
        spaces = defaultdict(lambda: {'used': set()})
        for p in points:
            m = re.match(r'%?([A-Z]+)(\d+)', p.address.upper())
            if m:
                spaces[m.group(1)]['used'].add(int(m.group(2)))
        if not spaces:
            return {}

        capacity_est = {
            'I': 65536, 'Q': 65536, 'M': 65536,
            'IW': 2048, 'QW': 2048, 'MW': 2048,
            'AI': 1024, 'AQ': 1024,
            'ID': 1024, 'QD': 1024, 'MD': 1024,
            'DB': 65536, 'V': 65536,
            'E': 65536, 'A': 65536,
            'X': 1024, 'Y': 1024,
        }

        result = {}
        for area, data in spaces.items():
            nums = sorted(data['used'])
            cap = capacity_est.get(area, 4096)
            util = round(len(nums) / cap * 100, 2)
            rng = f'{min(nums)} - {max(nums)}' if nums else 'N/A'
            if util > 80:
                assess = '⚠ 紧张，需规划扩展'
            elif util > 50:
                assess = '◉ 中等，可满足需求'
            else:
                assess = '✓ 充足'
            result[area] = {
                'count': len(nums),
                'range': rng,
                'capacity': cap,
                'utilization': util,
                'assessment': assess,
            }
        return result

    @staticmethod
    def _summarize_by_card(classified):
        """按槽位汇总"""
        summary = defaultdict(lambda: defaultdict(int))
        has_card = False
        for iotype, pts in classified.items():
            for p in pts:
                if p.card_slot:
                    has_card = True
                    summary[p.card_slot][iotype] += 1
        if not has_card:
            return {}
        for iotype, pts in classified.items():
            for p in pts:
                if not p.card_slot:
                    summary['未分配'][iotype] += 1
        return summary

    # ── CSV ──

    @staticmethod
    def _generate_csv(points, output_dir):
        classified = IOReportGenerator._classify(points)
        for iotype in IO_ORDER:
            pts = classified.get(iotype, [])
            cn = IO_LABELS.get(iotype, iotype)
            path = os.path.join(output_dir, f'IO_{iotype}.csv')
            with open(path, 'w', newline='', encoding='utf-8-sig') as f:
                w = csv.writer(f)
                w.writerow(['序号', '变量名', '地址', '数据类型', '备注', '槽位', '端子', '电缆'])
                for i, p in enumerate(pts, 1):
                    w.writerow([i, p.name, p.address, p.data_type, p.comment, p.card_slot, p.terminal, p.cable])

        full_path = os.path.join(output_dir, 'IO_点表总表.csv')
        with open(full_path, 'w', newline='', encoding='utf-8-sig') as f:
            w = csv.writer(f)
            w.writerow(['序号', '变量名', '地址', '数据类型', 'IO类型', '描述', '槽位', '端子', '电缆'])
            for i, p in enumerate(points, 1):
                cn = IO_LABELS.get(p.io_type, p.io_type)
                w.writerow([i, p.name, p.address, p.data_type, cn, p.comment, p.card_slot, p.terminal, p.cable])

    # ── Markdown ──

    @staticmethod
    def _generate_markdown(points, output_dir):
        classified = IOReportGenerator._classify(points)
        now = datetime.now().strftime('%Y-%m-%d %H:%M')

        lines = [
            '# IO 点表报告',
            '',
            f'- **生成时间:** {now}',
            f'- **总点数:** {len(points)}',
            '',
            '---',
            '',
            '## 1. 统计概览',
            '',
            '| 类型 | 说明 | 数量 | 占比 |',
            '|------|------|------|------|',
        ]

        for iotype in IO_ORDER:
            c = len(classified.get(iotype, []))
            pct = f'{c/len(points)*100:.1f}%' if points else '0%'
            lines.append(f'| **{iotype}** | {IO_LABELS.get(iotype, iotype)} | {c} | {pct} |')

        lines += ['', '### 分布图', '']
        max_c = max((len(classified.get(io, [])) for io in ['DI', 'DO', 'AI', 'AO']), default=1)
        for iotype in ['DI', 'DO', 'AI', 'AO']:
            c = len(classified.get(iotype, []))
            bl = int(c / max_c * 20) if max_c else 0
            bar = '▓' * bl + '░' * (20 - bl)
            lines.append(f'  {iotype:4s} |{bar}| {c}')

        lines += ['', '---', '', '## 2. IO 明细']

        for iotype in ['DI', 'DO', 'AI', 'AO']:
            pts = classified.get(iotype, [])
            if not pts:
                continue
            cn = IO_LABELS.get(iotype, iotype)
            lines += ['', f'### {iotype} - {cn} ({len(pts)})', '',
                      '| # | 变量名 | 地址 | 类型 | 槽位 | 备注 |',
                      '|---|--------|------|------|------|------|']
            for i, p in enumerate(pts, 1):
                lines.append(f'| {i} | {p.name} | {p.address} | {p.data_type} | {p.card_slot} | {p.comment} |')

        lines += ['', '---', '', '## 3. 地址占用']
        addr_info = IOReportGenerator._analyze_addresses(points)
        if addr_info:
            lines += ['', '| 区域 | 占用 | 范围 | 利用 | 评估 |',
                      '|------|------|------|------|------|']
            for area, info in sorted(addr_info.items()):
                lines.append(f'| %{area} | {info["count"]} | {info["range"]} | {info["utilization"]}% | {info["assessment"]} |')

        lines += ['', '---', '', '*报告由 HMI Toolbox 自动生成*', '']

        md_path = os.path.join(output_dir, 'IO点表报告.md')
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))

    # ── Word ──

    @staticmethod
    def generate_word(points, output_dir):
        """生成专业 Word 报告 — 封面 / 目录 / 统计 / 明细 / 地址分析 / 槽位 / 接线"""
        if not HAS_DOCX:
            print('  [WARN] python-docx 未安装，跳过 Word 导出')
            print('     安装: pip install python-docx')
            return None

        classified = IOReportGenerator._classify(points)
        now = datetime.now().strftime('%Y-%m-%d %H:%M')
        doc = Document()

        # 页边距
        for sec in doc.sections:
            sec.top_margin = Cm(2.0)
            sec.bottom_margin = Cm(2.0)
            sec.left_margin = Cm(2.5)
            sec.right_margin = Cm(2.5)

        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # ═══════════════════════════════════════
        #  封面
        # ═══════════════════════════════════════

        for _ in range(5):
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('IO 点表报告')
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('━' * 30)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('HMI 工程 IO 统计与地址分配明细')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        for _ in range(5):
            doc.add_paragraph()

        # 信息卡片（表格布局）
        info_data = [
            ('项目名称', 'HMI 工程 IO 统计'),
            ('生成时间', now),
            ('总 IO 点数', str(len(points))),
            ('文档版本', 'V2.0'),
        ]
        tbl_info = doc.add_table(rows=len(info_data), cols=2)
        tbl_info.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (lab, val) in enumerate(info_data):
            for ci, txt in enumerate([f'{lab}：', val]):
                cell = tbl_info.rows[i].cells[ci]
                cell.text = txt
                _set_word_cell_shading(cell, 'F5F7FA')
                if ci == 0:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for r in para.runs:
                            r.font.bold = True
                            r.font.size = Pt(10)
                            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        doc.add_page_break()

        # ═══════════════════════════════════════
        #  目录（手动）
        # ═══════════════════════════════════════

        _make_word_header(doc, '目录', 1)
        toc = [
            ('1', '统计概览'),
            ('2', 'IO 点明细'),
            ('  2.1', '数字量输入 (DI)'),
            ('  2.2', '数字量输出 (DO)'),
            ('  2.3', '模拟量输入 (AI)'),
            ('  2.4', '模拟量输出 (AO)'),
            ('3', '地址占用分析'),
            ('4', '模块/槽位分布'),
            ('5', '接线参考'),
        ]
        for num, title in toc:
            p = doc.add_paragraph()
            run = p.add_run(f'{num}  {title}')
            run.font.size = Pt(11)
            if not num.startswith(' '):
                run.font.bold = True

        doc.add_page_break()

        # ═══════════════════════════════════════
        #  1. 统计概览
        # ═══════════════════════════════════════

        _make_word_header(doc, '1. 统计概览', 1)
        _make_word_para(doc,
            '本报告按数字量输入 (DI)、数字量输出 (DO)、模拟量输入 (AI)、模拟量输出 (AO) 分类统计 IO 点位。')

        doc.add_paragraph()
        _make_word_header(doc, 'IO 类型统计表', 2)

        active_count = sum(len(classified.get(io, [])) for io in ['DI', 'DO', 'AI', 'AO'])

        headers_s = ['IO 类型', '中文说明', '数量', '占比']
        data_s = []
        for iotype in IO_ORDER:
            c = len(classified.get(iotype, []))
            pct = f'{c/len(points)*100:.1f}%' if points else '0%'
            data_s.append([iotype, IO_LABELS.get(iotype, iotype), str(c), pct])

        table_s = doc.add_table(rows=len(data_s) + 2, cols=4)
        table_s.alignment = WD_TABLE_ALIGNMENT.CENTER
        _style_word_table(table_s, headers_s, data_s)

        # 彩色 IO 类型标记
        for idx, (iotype, _, _, _) in enumerate(data_s, 1):
            if iotype in IO_COLORS_HEX:
                _set_word_cell_shading(table_s.rows[idx].cells[0], IO_COLORS_HEX[iotype][1:])
                for para in table_s.rows[idx].cells[0].paragraphs:
                    for r in para.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.bold = True

        # 合计行
        total_row = table_s.rows[-1]
        total_row.cells[0].text = '合计'
        total_row.cells[1].text = f'Active IO (DI/DO/AI/AO)'
        total_row.cells[2].text = str(active_count)
        total_row.cells[3].text = f'{active_count/len(points)*100:.1f}%' if points else '0%'
        for cell in total_row.cells:
            _set_word_cell_shading(cell, 'E8ECF1')
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.bold = True

        doc.add_paragraph()

        # 柱状图
        _make_word_header(doc, '数量分布图', 2)
        max_c = max((len(classified.get(io, [])) for io in ['DI', 'DO', 'AI', 'AO']), default=1)

        bar_tbl = doc.add_table(rows=5, cols=2)
        bar_tbl.style = 'Table Grid'
        for idx, iotype in enumerate(['DI', 'DO', 'AI', 'AO', '合计']):
            c = len(classified.get(iotype, [])) if iotype != '合计' else active_count
            cn = IO_LABELS.get(iotype, 'Active IO')
            label = f'{iotype} - {cn}' if iotype != '合计' else '合计 - Active IO'
            barlen = int(c / max_c * 25) if max_c else 0
            bar_chars = '█' * barlen
            cl = bar_tbl.rows[idx].cells[0]
            cl.width = Cm(5)
            cl.text = label
            cl2 = bar_tbl.rows[idx].cells[1]
            cl2.text = f'{bar_chars}{" " * (25 - barlen)}  {c}'
            if idx < 4 and iotype in IO_COLORS_HEX:
                _set_word_cell_shading(cl, IO_COLORS_HEX[iotype][1:])
                for para in cl.paragraphs:
                    for r in para.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.bold = True

        doc.add_page_break()

        # ═══════════════════════════════════════
        #  2. IO 明细
        # ═══════════════════════════════════════

        _make_word_header(doc, '2. IO 点明细', 1)

        for iotype in ['DI', 'DO', 'AI', 'AO']:
            pts = classified.get(iotype, [])
            if not pts:
                continue
            cn = IO_LABELS.get(iotype, iotype)
            _make_word_header(doc, f'{iotype} — {cn}（{len(pts)} 点）', 2)

            hdrs = ['#', '变量名', '地址', '数据类型', '槽位', '端子', '备注']
            cols_w = [Cm(1.0), Cm(4.5), Cm(2.8), Cm(2.2), Cm(1.8), Cm(1.8), Cm(3.5)]
            display_n = min(len(pts), 100)
            tbl = doc.add_table(rows=display_n + 1, cols=7)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, w in enumerate(cols_w):
                for row in tbl.rows:
                    row.cells[i].width = w

            _style_word_table(tbl, hdrs, [])

            for idx in range(display_n):
                p = pts[idx]
                row = tbl.rows[idx + 1]
                vals = [str(idx + 1), p.name, p.address, p.data_type, p.card_slot, p.terminal, p.comment]
                for ci, v in enumerate(vals):
                    row.cells[ci].text = v
                if idx % 2 == 0:
                    for ci in range(7):
                        _set_word_cell_shading(row.cells[ci], 'F5F7FA')

            if len(pts) > 100:
                _make_word_para(doc, f'… 以及 {len(pts) - 100} 个额外点位（详见 CSV 文件）',
                                size=9, color=RGBColor(0x99, 0x99, 0x99), italic=True)

            doc.add_paragraph()

        doc.add_page_break()

        # ═══════════════════════════════════════
        #  3. 地址占用分析
        # ═══════════════════════════════════════

        _make_word_header(doc, '3. 地址占用分析', 1)
        addr_info = IOReportGenerator._analyze_addresses(points)

        if addr_info:
            _make_word_para(doc, '各 PLC 地址区域的占用情况与扩展容量评估。')

            hdrs_a = ['地址区域', '占用点数', '地址范围', '利用率', '评估']
            data_a = []
            for area, info in sorted(addr_info.items()):
                data_a.append([
                    f'%{area}',
                    str(info['count']),
                    info['range'],
                    f'{info["utilization"]}%',
                    info['assessment'],
                ])
            tbl_a = doc.add_table(rows=len(data_a) + 1, cols=5)
            tbl_a.alignment = WD_TABLE_ALIGNMENT.CENTER
            _style_word_table(tbl_a, hdrs_a, data_a)

            # 利用率颜色标记
            for idx, (area, info) in enumerate(sorted(addr_info.items()), 1):
                u = info['utilization']
                if u > 80:
                    clr = 'F44336'
                elif u > 50:
                    clr = 'FF9800'
                else:
                    clr = '4CAF50'
                _set_word_cell_shading(tbl_a.rows[idx].cells[3], clr)
                for para in tbl_a.rows[idx].cells[3].paragraphs:
                    for r in para.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.bold = True

            doc.add_paragraph()
            total_u = sum(info['count'] for info in addr_info.values())
            _make_word_para(doc, f'已占用 {total_u} 个地址点，各区域均有余量。')
        else:
            _make_word_para(doc, '未检测到标准 PLC 地址格式，无法分析。',
                            color=RGBColor(0x99, 0x99, 0x99))

        doc.add_page_break()

        # ═══════════════════════════════════════
        #  4. 模块/槽位分布
        # ═══════════════════════════════════════

        _make_word_header(doc, '4. 模块 / 槽位分布', 1)
        card_summary = IOReportGenerator._summarize_by_card(classified)

        if card_summary:
            _make_word_para(doc, '各模块槽位的 IO 点分布，便于机柜布局与接线规划。')

            slots = sorted(card_summary.keys())
            hdrs_c = ['槽位', 'DI', 'DO', 'AI', 'AO', '小计']
            data_c = []
            total_pt = defaultdict(int)
            for slot in slots:
                counts = card_summary[slot]
                data_c.append([
                    slot,
                    str(counts.get('DI', 0)),
                    str(counts.get('DO', 0)),
                    str(counts.get('AI', 0)),
                    str(counts.get('AO', 0)),
                    str(sum(counts.values())),
                ])
                for t in ['DI', 'DO', 'AI', 'AO']:
                    total_pt[t] += counts.get(t, 0)

            tbl_c = doc.add_table(rows=len(data_c) + 2, cols=6)
            tbl_c.alignment = WD_TABLE_ALIGNMENT.CENTER
            _style_word_table(tbl_c, hdrs_c, data_c)

            # 合计行
            tr = tbl_c.rows[-1]
            tr.cells[0].text = '合计'
            for ci, t in enumerate(['DI', 'DO', 'AI', 'AO'], 1):
                tr.cells[ci].text = str(total_pt[t])
            tr.cells[5].text = str(sum(total_pt.values()))
            for cell in tr.cells:
                _set_word_cell_shading(cell, 'E8ECF1')
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.font.bold = True
        else:
            _make_word_para(doc, '槽位信息未填写。请确保 CSV 包含"槽位"列。',
                            color=RGBColor(0x99, 0x99, 0x99))

        doc.add_page_break()

        # ═══════════════════════════════════════
        #  5. 接线参考
        # ═══════════════════════════════════════

        _make_word_header(doc, '5. 接线参考', 1)

        wirings = [
            ['DI', '传感器 24V+ → DI 模块端子 → PLC 内部 → GND',
             'NPN/PNP 选型注意；使用屏蔽线'],
            ['DO', 'DO 模块端子 → 继电器线圈/KA → 24V+',
             '外接电源至公共端；感性负载加续流二极管'],
            ['AI', '传感器 Signal+ → AI 模块端子 → 模数转换',
             '屏蔽层单端接地；双绞屏蔽线'],
            ['AO', 'AO 模块端子 → 执行器 Signal+ → GND',
             '4-20mA 不掉线；频率输出注意阻抗匹配'],
        ]
        tbl_w = doc.add_table(rows=len(wirings) + 1, cols=3)
        tbl_w.alignment = WD_TABLE_ALIGNMENT.CENTER
        _style_word_table(tbl_w, ['类型', '接线路径', '注意事项'],
                          [[r[0], r[1], r[2]] for r in wirings])

        doc.add_paragraph()
        _make_word_header(doc, '接线示意图', 2)
        diagram = [
            '  传感器           DI 模块           PLC',
            '    ┌─┐             ┌───┐           ╔═══╗',
            '    │ │ 24V+ ───────┤ 1 ├───────────╢ I ║',
            '    │ │ COM  ───────┤ 2 ├───────────╢ 0 ║',
            '    └─┘             └───┘           ╚═══╝',
        ]
        for line in diagram:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(8)

        doc.add_paragraph()
        _make_word_para(doc, f'报告由 HMI Toolbox 自动生成 ({now})',
                        size=9, color=RGBColor(0xAA, 0xAA, 0xAA))
        _make_word_para(doc, '具体接线以电气原理图为准。',
                        size=9, color=RGBColor(0xAA, 0xAA, 0xAA), italic=True)

        word_path = os.path.join(output_dir, 'IO点表报告.docx')
        doc.save(word_path)
        print(f'  [OK] Word 报告: {word_path}')
        return word_path

    # ── PDF ──

    @staticmethod
    def generate_pdf(points, output_dir):
        """生成专业 PDF 报告 — 封面 / 统计 / 明细 / 地址 / 槽位 / 接线"""
        if not HAS_FPDF:
            print('  [WARN] fpdf2 未安装，跳过 PDF 导出')
            print('     安装: pip install fpdf2')
            return None

        classified = IOReportGenerator._classify(points)
        now = datetime.now().strftime('%Y-%m-%d %H:%M')
        font_path = find_chinese_font()

        class PDFReport(FPDF):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self.fp = font_path
                self.set_auto_page_break(auto=True, margin=22)
                if self.fp:
                    self.add_font('CJK', '', self.fp)
                    self.add_font('CJK', 'B', self.fp)

            def _cf(self, style='', size=10):
                return ('CJK' if self.fp else 'Helvetica', style, size)

            def header(self):
                if self.page_no() <= 1:
                    return
                self.set_font(*self._cf('', 8))
                self.set_text_color(153, 153, 153)
                self.cell(0, 8, 'IO 点表报告', align='L')
                self.cell(0, 8, f'第 {self.page_no()} 页', align='R', new_x='LMARGIN', new_y='NEXT')
                self.set_draw_color(200, 200, 200)
                self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
                self.ln(5)

            def footer(self):
                self.set_y(-15)
                self.set_font(*self._cf('', 7))
                self.set_text_color(180, 180, 180)
                self.cell(0, 10, f'HMI Toolbox  |  {now}', align='C')

            def heading(self, text, size=16):
                self.ln(4)
                self.set_font(*self._cf('B', size))
                self.set_text_color(15, 23, 42)
                self.cell(0, size * 0.45, text, new_x='LMARGIN', new_y='NEXT')
                self.ln(4)

            def body(self, text, size=9, color=(51, 51, 51)):
                self.set_font(*self._cf('', size))
                self.set_text_color(*color)
                self.multi_cell(0, size * 0.5, text)

            def draw_table(self, headers, data, col_widths, header_color=(15, 23, 42)):
                total_w = sum(col_widths)

                # Header
                self.set_font(*self._cf('B', 9))
                self.set_fill_color(*header_color)
                self.set_text_color(255, 255, 255)
                for i, h in enumerate(headers):
                    self.cell(col_widths[i], 10, f' {h}', border=0, fill=True, align='C')
                self.ln()

                # Data
                self.set_font(*self._cf('', 8))
                fill = False
                for row_data in data:
                    if self.get_y() + 10 > self.h - 25:
                        self.add_page()
                        self.set_font(*self._cf('B', 9))
                        self.set_fill_color(*header_color)
                        self.set_text_color(255, 255, 255)
                        for i, h in enumerate(headers):
                            self.cell(col_widths[i], 10, f' {h}', border=0, fill=True, align='C')
                        self.ln()
                        self.set_font(*self._cf('', 8))

                    self.set_fill_color(245, 247, 250) if fill else self.set_fill_color(255, 255, 255)
                    self.set_text_color(51, 51, 51)
                    for i, cell_val in enumerate(row_data):
                        text = str(cell_val)[:40]
                        self.cell(col_widths[i], 9, f' {text}', border=0, fill=True,
                                  align='L' if i > 0 else 'C')
                    self.ln()
                    fill = not fill

        pdf = PDFReport()
        pdf.set_margins(20, 20, 20)
        pdf.add_page()

        # ═══════════ 封面 ═══════════

        pdf.ln(60)
        pdf.set_font(*pdf._cf('B', 32))
        pdf.set_text_color(15, 23, 42)
        pdf.cell(0, 20, 'IO 点表报告', align='C', new_x='LMARGIN', new_y='NEXT')
        pdf.set_font(*pdf._cf('', 14))
        pdf.set_text_color(102, 102, 102)
        pdf.cell(0, 12, 'HMI 工程 IO 统计与地址分配明细', align='C', new_x='LMARGIN', new_y='NEXT')
        pdf.ln(15)
        pdf.set_draw_color(26, 86, 219)
        pdf.set_line_width(0.5)
        pdf.line(60, pdf.get_y(), pdf.w - 60, pdf.get_y())
        pdf.ln(15)

        for label, value in [('项目名称', 'HMI 工程 IO 统计'),
                             ('生成时间', now),
                             ('总 IO 点数', str(len(points)))]:
            pdf.set_font(*pdf._cf('', 10))
            pdf.set_text_color(15, 23, 42)
            pdf.cell(35, 10, label, align='R')
            pdf.set_text_color(51, 51, 51)
            pdf.cell(0, 10, f'：{value}', align='L', new_x='LMARGIN', new_y='NEXT')

        pdf.add_page()

        # ═══════════ 1. 统计 ═══════════

        pdf.heading('1. 统计概览', 18)
        pdf.body('按数字量输入 (DI)、数字量输出 (DO)、模拟量输入 (AI)、模拟量输出 (AO) 分类统计。')
        pdf.ln(5)

        active_count = sum(len(classified.get(io, [])) for io in ['DI', 'DO', 'AI', 'AO'])
        pdf.heading('IO 类型统计表', 13)

        hdrs = ['IO 类型', '中文说明', '数量', '占比']
        cw = [22, 55, 22, 22]
        data = []
        for iotype in IO_ORDER:
            c = len(classified.get(iotype, []))
            pct = f'{c/len(points)*100:.1f}%' if points else '0%'
            data.append([iotype, IO_LABELS.get(iotype, iotype), str(c), pct])
        pdf.draw_table(hdrs, data, cw)
        pdf.ln(3)

        # 合计
        pdf.set_font(*pdf._cf('B', 9))
        pdf.set_fill_color(232, 236, 241)
        pdf.set_text_color(51, 51, 51)
        pdf.cell(22, 9, ' 合计', fill=True)
        pdf.cell(55, 9, ' Active IO', fill=True)
        pdf.cell(22, 9, f' {active_count}', fill=True)
        pct_all = f'{active_count/len(points)*100:.1f}%' if points else '0%'
        pdf.cell(22, 9, f' {pct_all}', fill=True)
        pdf.ln(8)

        # Bar chart
        pdf.heading('数量分布图', 13)
        max_c = max((len(classified.get(io, [])) for io in ['DI', 'DO', 'AI', 'AO']), default=1)
        for iotype in ['DI', 'DO', 'AI', 'AO']:
            c = len(classified.get(iotype, []))
            bw = max(int(c / max_c * 80) if max_c else 10, 8)
            cn = IO_LABELS.get(iotype, iotype)
            pdf.set_font(*pdf._cf('B', 9))
            pdf.cell(35, 8, f' {iotype} - {cn}', align='R')
            pdf.set_fill_color(*IO_COLORS_RGB[iotype])
            pdf.cell(bw, 8, '', fill=True)
            pdf.set_font(*pdf._cf('', 9))
            pdf.set_text_color(51, 51, 51)
            pdf.cell(5)
            pdf.cell(20, 8, f'{c}', new_x='LMARGIN', new_y='NEXT')
            pdf.ln(2)

        pdf.add_page()

        # ═══════════ 2. IO 明细 ═══════════

        pdf.heading('2. IO 点明细', 18)

        for iotype in ['DI', 'DO', 'AI', 'AO']:
            pts = classified.get(iotype, [])
            if not pts:
                continue
            cn = IO_LABELS.get(iotype, iotype)
            pdf.heading(f'{iotype} — {cn}（{len(pts)} 点）', 13)

            display = pts[:50]
            hdrs = ['#', '变量名', '地址', '类型', '备注']
            cw = [10, 48, 30, 22, 45]
            data = [[str(i + 1), p.name, p.address, p.data_type, p.comment]
                    for i, p in enumerate(display)]
            pdf.draw_table(hdrs, data, cw)
            pdf.ln(2)

            if len(pts) > 50:
                pdf.body(f'… 以及 {len(pts) - 50} 个额外点位（详见 CSV 文件）', 7, (153, 153, 153))

        pdf.add_page()

        # ═══════════ 3. 地址分析 ═══════════

        pdf.heading('3. 地址占用分析', 18)
        addr_info = IOReportGenerator._analyze_addresses(points)
        if addr_info:
            hdrs = ['地址区域', '占用', '范围', '利用率']
            cw = [28, 20, 42, 30]
            data = []
            for area, info in sorted(addr_info.items()):
                data.append([f'%{area}', str(info['count']), info['range'],
                             f'{info["utilization"]}%'])
            pdf.draw_table(hdrs, data, cw)
            pdf.ln(4)
            total_u = sum(info['count'] for info in addr_info.values())
            pdf.body(f'已占用 {total_u} 个地址点，各区域余量充足。')
        else:
            pdf.body('未检测到标准 PLC 地址格式。', 9, (153, 153, 153))

        pdf.ln(6)

        # ═══════════ 4. 槽位分布 ═══════════

        pdf.heading('4. 模块 / 槽位分布', 18)
        card_summary = IOReportGenerator._summarize_by_card(classified)
        if card_summary:
            hdrs = ['槽位', 'DI', 'DO', 'AI', 'AO', '小计']
            cw = [30, 18, 18, 18, 18, 22]
            data = []
            tp = defaultdict(int)
            for slot in sorted(card_summary.keys()):
                cnt = card_summary[slot]
                data.append([slot, str(cnt.get('DI', 0)), str(cnt.get('DO', 0)),
                             str(cnt.get('AI', 0)), str(cnt.get('AO', 0)),
                             str(sum(cnt.values()))])
                for t in ['DI', 'DO', 'AI', 'AO']:
                    tp[t] += cnt.get(t, 0)
            pdf.draw_table(hdrs, data, cw)
        else:
            pdf.body('槽位信息未填写。', 9, (153, 153, 153))

        pdf.ln(6)

        # ═══════════ 5. 接线 ═══════════

        pdf.heading('5. 接线参考', 18)
        wirings = [
            ['DI', '传感器 24V+ → DI 模块端子 → PLC → GND'],
            ['DO', 'DO 模块端子 → 继电器线圈 → 24V+'],
            ['AI', '传感器 Signal+ → AI 模块端子 → Shield/GND'],
            ['AO', 'AO 模块端子 → 执行器 Signal+ → GND'],
        ]
        pdf.draw_table(['类型', '接线路径'], wirings, [25, 130])
        pdf.ln(4)
        pdf.body('本报告供 HMI 开发和现场调试参考，具体接线以电气原理图为准。',
                 8, (153, 153, 153))

        pdf_path = os.path.join(output_dir, 'IO点表报告.pdf')
        pdf.output(pdf_path)
        print(f'  [OK] PDF 报告: {pdf_path}')
        return pdf_path


# ═══════════════════════════════════════════════════════════════
# 总调度
# ═══════════════════════════════════════════════════════════════

def generate_all(points, output_dir, args=None):
    """生成所有格式的报告"""
    if args is None:
        class A: skip_csv = False; skip_word = False; skip_pdf = False; skip_markdown = False
        args = A()

    os.makedirs(output_dir, exist_ok=True)

    if not args.skip_csv:
        print('[DATA] 生成 CSV 报表...')
        IOReportGenerator._generate_csv(points, output_dir)
        print(f'  [OK] IO_点表总表.csv — {len(points)} 个点')
        print()

    if not args.skip_markdown:
        print('[gen] 生成 Markdown 报告...')
        IOReportGenerator._generate_markdown(points, output_dir)
        print('  [OK] IO点表报告.md')
        print()

    if not args.skip_word:
        print('[DOC] 生成 Word 报告...')
        IOReportGenerator.generate_word(points, output_dir)
        print()

    if not args.skip_pdf:
        print('[PDF] 生成 PDF 报告...')
        IOReportGenerator.generate_pdf(points, output_dir)
        print()

    print('─' * 50)
    print('输出文件:')
    print(f'  [DATA]  {output_dir}\\IO_点表总表.csv')
    print(f'  [DATA]  {output_dir}\\IO_DI.csv 等（分类型）')
    print(f'  [gen]  {output_dir}\\IO点表报告.md')
    if HAS_DOCX and not (args and args.skip_word):
        print(f'  [DOC]  {output_dir}\\IO点表报告.docx')
    if HAS_FPDF and not (args and args.skip_pdf):
        print(f'  [PDF]  {output_dir}\\IO点表报告.pdf')

# ═══════════════════════════════════════════════════════════════
# 示例数据
# ═══════════════════════════════════════════════════════════════

def generate_example(output_dir):
    os.makedirs(output_dir, exist_ok=True)
    examples = [
        IOPoint('急停按钮', '%I0.0', 'BOOL', '急停 E-Stop', 'DI', '0', '1', 'DI_8W_01'),
        IOPoint('启动按钮', '%I0.1', 'BOOL', '启动 Start', 'DI', '0', '2', 'DI_8W_01'),
        IOPoint('停止按钮', '%I0.2', 'BOOL', '停止 Stop', 'DI', '0', '3', 'DI_8W_01'),
        IOPoint('复位按钮', '%I0.3', 'BOOL', '故障复位 Reset', 'DI', '0', '4', 'DI_8W_01'),
        IOPoint('光电传感器1', '%I1.0', 'BOOL', '工件检测', 'DI', '1', '1', 'DI_8W_02'),
        IOPoint('光电传感器2', '%I1.1', 'BOOL', '到位检测', 'DI', '1', '2', 'DI_8W_02'),
        IOPoint('限位开关_左', '%I2.0', 'BOOL', '左限位', 'DI', '2', '1', 'DI_8W_03'),
        IOPoint('限位开关_右', '%I2.1', 'BOOL', '右限位', 'DI', '2', '2', 'DI_8W_03'),
        IOPoint('温度开关', '%I3.0', 'BOOL', '超温报警', 'DI', '3', '1', 'DI_8W_04'),
        IOPoint('压力开关', '%I3.1', 'BOOL', '欠压报警', 'DI', '3', '2', 'DI_8W_04'),
        IOPoint('安全光幕', '%I4.0', 'BOOL', '安全光幕', 'DI', '3', '3', 'DI_8W_04'),
        IOPoint('编码器Z脉冲', '%I5.0', 'BOOL', '编码器零点', 'DI', '4', '1', 'DI_8W_05'),
        IOPoint('电机运行', '%Q0.0', 'BOOL', '主电机', 'DO', '5', '1', 'DO_8W_01'),
        IOPoint('电机正转', '%Q0.1', 'BOOL', '正转接触器', 'DO', '5', '2', 'DO_8W_01'),
        IOPoint('电机反转', '%Q0.2', 'BOOL', '反转接触器', 'DO', '5', '3', 'DO_8W_01'),
        IOPoint('报警指示灯', '%Q1.0', 'BOOL', '红色指示灯', 'DO', '6', '1', 'DO_8W_02'),
        IOPoint('运行指示灯', '%Q1.1', 'BOOL', '绿色指示灯', 'DO', '6', '2', 'DO_8W_02'),
        IOPoint('蜂鸣器', '%Q1.2', 'BOOL', '报警蜂鸣', 'DO', '6', '3', 'DO_8W_02'),
        IOPoint('电磁阀1', '%Q2.0', 'BOOL', '进料阀', 'DO', '7', '1', 'DO_8W_03'),
        IOPoint('电磁阀2', '%Q2.1', 'BOOL', '排料阀', 'DO', '7', '2', 'DO_8W_03'),
        IOPoint('冷却风扇', '%Q3.0', 'BOOL', '散热风扇', 'DO', '7', '3', 'DO_8W_03'),
        IOPoint('加热器', '%Q3.1', 'BOOL', '加热器通断', 'DO', '7', '4', 'DO_8W_03'),
        IOPoint('温度传感器1', '%AIW0', 'INT', '炉温1 0-200°C', 'AI', '8', '1', 'AI_4W_01'),
        IOPoint('温度传感器2', '%AIW2', 'INT', '炉温2 0-200°C', 'AI', '8', '2', 'AI_4W_01'),
        IOPoint('压力变送器', '%AIW4', 'INT', '系统压力 0-16bar', 'AI', '8', '3', 'AI_4W_02'),
        IOPoint('流量计', '%AIW6', 'INT', '瞬时流量', 'AI', '9', '1', 'AI_2W_01'),
        IOPoint('液位计', '%AIW8', 'INT', '液位 0-500mm', 'AI', '9', '2', 'AI_2W_01'),
        IOPoint('变频器频率', '%AQW0', 'INT', '变频器给定 0-50Hz', 'AO', '10', '1', 'AO_01'),
        IOPoint('调节阀开度', '%AQW2', 'INT', '阀门控制 0-100%', 'AO', '10', '2', 'AO_01'),
    ]

    csv_path = os.path.join(output_dir, 'tags_input.csv')
    with open(csv_path, 'w', newline='', encoding='utf-8-sig') as f:
        w = csv.writer(f)
        w.writerow(['变量名', '地址', '数据类型', '备注', '槽位', '端子', '电缆'])
        for p in examples:
            w.writerow([p.name, p.address, p.data_type, p.comment, p.card_slot, p.terminal, p.cable])
    return examples


# ═══════════════════════════════════════════════════════════════
# 入口
# ═══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description='IO 点表自动生成报告 v2.0 — 支持 Word · PDF · Markdown · CSV',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''示例:
  python io_report_generator.py --example
  python io_report_generator.py -i tags.csv -o output_dir
  python io_report_generator.py -i tags.csv --skip-word --skip-pdf
''')
    parser.add_argument('-i', '--input', help='输入 CSV 标签文件')
    parser.add_argument('-o', '--output', default='./output', help='输出目录 (默认: ./output)')
    parser.add_argument('--example', action='store_true', help='生成示例报告')
    parser.add_argument('--skip-csv', action='store_true', help='跳过 CSV 导出')
    parser.add_argument('--skip-word', action='store_true', help='跳过 Word 导出')
    parser.add_argument('--skip-pdf', action='store_true', help='跳过 PDF 导出')
    parser.add_argument('--skip-markdown', action='store_true', help='跳过 Markdown 导出')

    args = parser.parse_args()
    output_dir = os.path.abspath(args.output)

    print('=' * 56)
    print('  [OUT] IO 点表报告生成器 v2.0')
    print('  输出: Word · PDF · Markdown · CSV')
    print('=' * 56)
    print()

    if args.example:
        print('[gen] 生成示例报告…')
        os.makedirs(output_dir, exist_ok=True)
        points = generate_example(output_dir)
        print(f'  - 输入: {os.path.join(output_dir, "tags_input.csv")}')
        print(f'  - 解析: {len(points)} 个 IO 点')
        print()
        generate_all(points, output_dir, args)
        print(f'\n[OK] 完成 → {output_dir}')
        return

    if not args.input:
        parser.print_help()
        print()
        print('💡 试试: python io_report_generator.py --example')
        sys.exit(1)

    if not os.path.exists(args.input):
        print(f'[ERR] 文件不存在: {args.input}')
        sys.exit(1)

    os.makedirs(output_dir, exist_ok=True)
    print(f'📂 输入: {args.input}')
    points = IOReportGenerator.read_tags(args.input)
    if not points:
        print('  [WARN] 未解析到有效 IO 点')
        print('  预期列名: 变量名, 地址, 数据类型, 备注')
        sys.exit(1)

    print(f'  - 解析: {len(points)} 个 IO 点')
    print()
    generate_all(points, output_dir, args)
    print(f'\n[OK] 完成 → {output_dir}')


if __name__ == '__main__':
    main()
